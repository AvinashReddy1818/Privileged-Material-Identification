from flask import Flask, request, render_template, send_from_directory, redirect, url_for, session
import os
import time
from werkzeug.utils import secure_filename
from PyPDF2 import PdfWriter, PdfReader
from docx import Document
import fitz  # PyMuPDF
import re
import mysql.connector
from datetime import timedelta

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}

app.secret_key = 'avin'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(minutes=30)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def generate_unique_filename(base_name, extension):
    timestamp = time.strftime("%Y%m%d%H%M%S")
    return f"{base_name}_{timestamp}.{extension}"

def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        text = ' '.join([para.text for para in doc.paragraphs])
        return text, doc
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return "", None

def extract_text_from_pdf(file_path):
    try:
        reader = PdfReader(file_path)
        text = ''.join(page.extract_text() or "" for page in reader.pages)
        return text, reader
    except Exception as e:
        print(f"Error extracting PDF: {e}")
        return "", None

def extract_text_from_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read(), None
    except Exception as e:
        print(f"Error extracting TXT: {e}")
        return "", None

def highlight_in_docx(doc, keywords):
    for para in doc.paragraphs:
        for run in para.runs:
            for keyword in keywords:
                if keyword.lower() in run.text.lower():
                    highlighted_text = re.sub(rf'(\b{re.escape(keyword)}\b)', r'[\1]', run.text, flags=re.IGNORECASE)
                    run.text = highlighted_text
                    run.font.highlight_color = 3
    return doc

def highlight_in_pdf(file_path, keywords, output_path):
    try:
        doc = fitz.open(file_path)
        for page in doc:
            for keyword in keywords:
                areas = page.search_for(keyword)
                for area in areas:
                    highlight = page.add_highlight_annot(area)
                    highlight.update()
        doc.save(output_path)
    except Exception as e:
        print(f"Error highlighting PDF: {e}")

def highlight_in_txt(text, keywords):
    for keyword in keywords:
        text = re.sub(rf'(\b{re.escape(keyword)}\b)', r'[\1]', text, flags=re.IGNORECASE)
    return text

def convert_txt_to_pdf(text, output_path):
    pdf_writer = PdfWriter()
    page = fitz.open()
    page.insert_text((50, 50), text)
    pdf_writer.add_page(page)
    with open(output_path, 'wb') as f:
        pdf_writer.write(f)

def get_db_connection():
    return mysql.connector.connect(
        host='localhost',
        user='root',
        password='',
        database='attorney'
    )

@app.route('/')
def index():
    return render_template('login.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        connection = get_db_connection()
        cursor = connection.cursor(dictionary=True)
        cursor.execute('SELECT * FROM users WHERE username = %s', (username,))
        user = cursor.fetchone()
        connection.close()
        if user and user['password'] == password:
            session['user_id'] = user['id']
            session['username'] = user['username']
            return render_template('dashboard.html',username=username)
        return 'Invalid credentials'


@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        connection = get_db_connection()
        cursor = connection.cursor()
        cursor.execute('INSERT INTO users (username, password) VALUES (%s, %s)', (username, password))
        connection.commit()
        connection.close()
        return render_template('login.html',username=username)

@app.route('/logout')
def logout():
    session.pop('user_id', None)
    session.pop('username', None)
    return redirect(url_for('index'))

@app.route('/dashboard', methods=['GET', 'POST'])
def dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    original_pdf_filename = None
    highlighted_pdf_filename = None

    if request.method == 'POST':
        file = request.files.get('file')
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)

            keywords = request.form.get('keywords', '')
            keywords_list = [k.strip() for k in keywords.split(',') if k.strip()]

            output_filename = generate_unique_filename(filename.rsplit('.', 1)[0], 'pdf')
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

            if file_path.endswith('.docx'):
                text, doc = extract_text_from_docx(file_path)
                if doc:
                    doc = highlight_in_docx(doc, keywords_list)
                    temp_output_path = os.path.join(app.config['OUTPUT_FOLDER'], generate_unique_filename(filename.rsplit('.', 1)[0], 'docx'))
                    doc.save(temp_output_path)
                    convert_txt_to_pdf(open(temp_output_path).read(), output_path)
                    os.remove(temp_output_path)

            elif file_path.endswith('.pdf'):
                text, _ = extract_text_from_pdf(file_path)
                if text:
                    highlight_in_pdf(file_path, keywords_list, output_path)

            elif file_path.endswith('.txt'):
                text, _ = extract_text_from_txt(file_path)
                if text:
                    highlighted_text = highlight_in_txt(text, keywords_list)
                    convert_txt_to_pdf(highlighted_text, output_path)

            else:
                return "Unsupported file format"

            # Save file metadata to the database
            connection = get_db_connection()
            cursor = connection.cursor()
            cursor.execute('INSERT INTO file_metadata (user_id, original_filename, output_filename, keywords, created_at) VALUES (%s, %s, %s, %s, NOW())',
                           (session['user_id'], filename, output_filename, keywords))
            connection.commit()
            connection.close()

            original_pdf_filename = filename
            highlighted_pdf_filename = output_filename

    return render_template('dashboard.html', 
                           username=session.get('username'),
                           original_pdf_filename=original_pdf_filename,
                           highlighted_pdf_filename=highlighted_pdf_filename)



@app.route('/activity')
def activity():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    connection = get_db_connection()
    cursor = connection.cursor(dictionary=True)
    cursor.execute('SELECT * FROM file_metadata WHERE user_id = %s ORDER BY created_at DESC', (session['user_id'],))
    files = cursor.fetchall()
    connection.close()
    return render_template('activity.html', files=files)

@app.route('/outputs/<filename>')
def download_file(filename):
    return send_from_directory(app.config['OUTPUT_FOLDER'], filename)

@app.route('/uploads/<filename>')
def upload_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

if __name__ == '__main__':
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
    if not os.path.exists(app.config['OUTPUT_FOLDER']):
        os.makedirs(app.config['OUTPUT_FOLDER'])
    app.run(debug=True)